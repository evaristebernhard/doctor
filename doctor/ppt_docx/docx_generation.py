'''将大模型生成的json数据转换为word，可修改代码自定义word的样式'''
# 导入标准库
import os  # 操作系统接口模块，用于文件路径操作
import re  # 正则表达式模块，用于文本处理
import hashlib  # 哈希算法模块，用于生成文件名
import time  # 时间模块，用于生成时间戳

# 导入第三方库
from docx import Document  # Python-Docx库，用于创建和操作Word文档
from docx.shared import Pt  # 用于设置字体大小的单位（磅）
from docx.oxml.ns import qn  # 用于处理Word XML命名空间，特别是中文字体支持
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # 用于设置段落对齐方式
from typing import Dict  # 类型提示，用于字典类型注解

# 导入项目模块
from env import get_app_root  # 获取应用根目录的函数

# 定义Word文件输出目录路径，将Word文件保存在应用根目录下的data/cache/docx文件夹中
_OUTPUT_DIR_DOCX = os.path.join(get_app_root(), "data/cache/docx")

# 如果文件夹路径不存在，先创建该目录
if not os.path.exists(_OUTPUT_DIR_DOCX):
    # 递归创建目录，确保所有父目录都存在
    os.makedirs(_OUTPUT_DIR_DOCX)

def get_file_path_docx(text):
    """
    生成唯一的Word文件路径
    
    Args:
        text (str): 用于生成文件名的文本
        
    Returns:
        str: 完整的文件路径
    """
    # 使用SHA256哈希算法对文本进行哈希，生成唯一的文件名
    file_name = hashlib.sha256(text.encode("utf-8")).hexdigest()  # 可以使用uuid替代
    # 返回完整的文件路径，文件扩展名为.docx
    return os.path.join(_OUTPUT_DIR_DOCX, f"{file_name}.docx")

def is_chinese(text: str) -> bool:
    """
    判断文本中是否包含中文字符
    
    Args:
        text (str): 需要检测的文本
        
    Returns:
        bool: 如果包含中文字符返回True，否则返回False
    """
    # 使用正则表达式检测是否包含中文字符（Unicode范围\u4e00-\u9fff）
    return bool(re.search(r'[\u4e00-\u9fff]', text))

def generate_docx_content(docx_content: Dict) -> str:
    """
    根据内容字典生成Word文档文件
    
    Args:
        docx_content (Dict): 包含Word内容的字典，应包含标题和章节内容
        
    Returns:
        str: 生成的Word文件路径
    """
    # 创建一个新的Word文档对象
    document = Document()

    # Word 标题 - 添加文档标题（级别0表示标题）
    title_heading = document.add_heading(docx_content['title'], 0)
    # 设置标题居中对齐
    title_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # 获取标题段落的第一个运行（文本块）
    title_run = title_heading.runs[0]
    
    # 根据标题是否包含中文设置字体
    if is_chinese(docx_content['title']):
        # 如果包含中文，设置中文字体为黑体
        title_run.font.name = '黑体'  # 中文字体
        # 通过XML设置中文字体属性
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')  # 设置中文字体
    else:
        # 如果不包含中文，设置西文字体为Arial
        title_run.font.name = 'Arial'  
    
    # 设置标题字体大小为24磅
    title_run.font.size = Pt(24)  # 标题字体大小

    # 页内容 - 遍历所有章节内容
    print(f'总共 {len(docx_content["sections"])} 个章节')
    # 遍历sections数组中的每个章节
    for i, section in enumerate(docx_content['sections']):
        # 打印当前章节信息到控制台
        print(f'生成第 {i + 1} 章节: {section["heading"]}')
        # 添加章节标题（级别1表示一级标题）
        section_heading = document.add_heading(section['heading'], level=1)
        # 获取章节标题段落的第一个运行
        section_heading_run = section_heading.runs[0]
        
        # 根据章节标题是否包含中文设置字体
        if is_chinese(section['heading']):
            # 如果包含中文，设置中文字体为宋体
            section_heading_run.font.name = '宋体'
            # 通过XML设置中文字体属性
            section_heading_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        else:
            # 如果不包含中文，设置西文字体为Times New Roman
            section_heading_run.font.name = 'Times New Roman'
        
        # 设置章节标题字体大小为16磅
        section_heading_run.font.size = Pt(16)  # 章节标题字体大小

        # 遍历当前章节的所有段落
        for paragraph in section['paragraphs']:
            # 添加段落标题（级别2表示二级标题）
            para_heading = document.add_heading(paragraph['heading'], level=2)
            # 获取段落标题的第一个运行
            para_heading_run = para_heading.runs[0]
            
            # 根据段落标题是否包含中文设置字体
            if is_chinese(paragraph['heading']):
                # 如果包含中文，设置中文字体为宋体
                para_heading_run.font.name = '宋体'
                # 通过XML设置中文字体属性
                para_heading_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            else:
                # 如果不包含中文，设置西文字体为Calibri
                para_heading_run.font.name = 'Calibri'
            
            # 设置段落标题字体大小为14磅
            para_heading_run.font.size = Pt(14)  # 段落标题字体大小

            # 设置正文内容字体 - 添加段落正文内容
            p = document.add_paragraph(paragraph['content'])
            # 获取正文段落的第一个运行
            p_run = p.runs[0]
            
            # 根据正文内容是否包含中文设置字体
            if is_chinese(paragraph['content']):
                # 如果包含中文，设置中文字体为宋体
                p_run.font.name = '宋体'
                # 通过XML设置中文字体属性
                p_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            else:
                # 如果不包含中文，设置西文字体为Arial
                p_run.font.name = 'Arial'
            
            # 设置正文字体大小为12磅
            p_run.font.size = Pt(12)  # 正文字体大小

    # 生成输出文件路径，使用当前时间戳确保文件名唯一
    _output_file = get_file_path_docx(str(time.time()))
    # 保存Word文档到指定路径
    document.save(_output_file)

    # 返回生成的Word文件路径
    return _output_file